#!/usr/bin/env python
import argparse
from html import escape
import json
from pathlib import Path

try:
    from docx import Document  # type: ignore
    from docx.enum.style import WD_STYLE_TYPE  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from reportlab.lib.enums import TA_CENTER  # type: ignore
    from reportlab.lib.pagesizes import A4  # type: ignore
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore
    from reportlab.lib.units import mm  # type: ignore
    from reportlab.pdfbase import pdfmetrics  # type: ignore
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore
except ImportError as exc:
    raise RuntimeError(
        "缺少 portable Python 依赖 python-docx、lxml 或 reportlab。请校验 runtime/python 并重新运行 deploy.ps1。"
    ) from exc

try:
    from PIL import Image  # type: ignore
except Exception:
    Image = None


ACCENT = RGBColor(0x13, 0x5D, 0x66)
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def normalize_text(value):
    return str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()


def ensure_styles(doc: Document) -> None:
    styles = doc.styles
    if "Article Title" not in styles:
        style = styles.add_style("Article Title", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(24)
        style.font.bold = True
        style.font.color.rgb = TEXT
    if "Article Lead" not in styles:
        style = styles.add_style("Article Lead", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11.5)
        style.font.color.rgb = MUTED
    if "Section Heading CN" not in styles:
        style = styles.add_style("Section Heading CN", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(15)
        style.font.bold = True
        style.font.color.rgb = ACCENT
    if "Article Body CN" not in styles:
        style = styles.add_style("Article Body CN", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11.5)
        style.font.color.rgb = TEXT


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.6):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_image_if_exists(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        return
    width = Inches(5.7)
    if Image is not None:
        try:
            with Image.open(image_path) as image:
                img_w, img_h = image.size
                if img_w and img_h and img_h > img_w:
                    width = Inches(4.2)
        except Exception:
            pass
    doc.add_picture(str(image_path), width=width)
    caption_p = doc.add_paragraph(style="Article Lead")
    caption_run = caption_p.add_run(caption)
    caption_run.italic = True
    set_paragraph_spacing(caption_p, after=10, line=1.2)


def build_docx(package: dict, output_dir: Path) -> Path:
    doc = Document()
    ensure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    hero = doc.add_table(rows=1, cols=1)
    hero.autofit = True
    hero_cell = hero.rows[0].cells[0]
    set_cell_shading(hero_cell, "F3F8F8")
    title_p = hero_cell.paragraphs[0]
    title_p.style = doc.styles["Article Title"]
    title_p.add_run(normalize_text(package.get("title", "微信公众号文章")))
    set_paragraph_spacing(title_p, after=6, line=1.1)

    summary_p = hero_cell.add_paragraph(style="Article Lead")
    summary_p.add_run(normalize_text(package.get("summary", "")))
    set_paragraph_spacing(summary_p, after=0, line=1.45)

    assets = package.get("asset_images", []) or []
    assets_dir = output_dir / "extracted-paper-assets"

    if assets:
        doc.add_paragraph("")
        add_image_if_exists(doc, assets_dir / assets[0], "论文中提取的代表性配图，可直接用于封面或导语之后。")

    for index, section_data in enumerate(package.get("sections", []), start=1):
        heading = doc.add_paragraph(style="Section Heading CN")
        heading.add_run(normalize_text(section_data.get("heading", f"第 {index} 部分")))
        set_paragraph_spacing(heading, before=10, after=6, line=1.2)

        for paragraph in section_data.get("paragraphs", []):
            body = doc.add_paragraph(style="Article Body CN")
            body.add_run(normalize_text(paragraph))
            set_paragraph_spacing(body, after=6)

        if index < len(assets):
            add_image_if_exists(
                doc,
                assets_dir / assets[index],
                f"论文配图 {index + 1}，建议在正文对应段落附近穿插使用。",
            )

    note_heading = doc.add_paragraph(style="Section Heading CN")
    note_heading.add_run("发布前检查")
    set_paragraph_spacing(note_heading, before=8, after=6, line=1.2)
    for step in package.get("next_steps", []):
        bullet = doc.add_paragraph(style="Article Body CN")
        bullet.add_run("• " + normalize_text(step))
        set_paragraph_spacing(bullet, after=4)

    docx_path = output_dir / "popular-article.docx"
    doc.save(str(docx_path))
    return docx_path


def build_pdf(package: dict, output_dir: Path) -> Path:
    pdf_path = output_dir / "popular-article.pdf"
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=20,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=10 * mm,
    )
    lead_style = ParagraphStyle(
        "ChineseLead",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=18,
        textColor="#6B7280",
        spaceAfter=7 * mm,
    )
    heading_style = ParagraphStyle(
        "ChineseHeading",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=15,
        leading=22,
        textColor="#135D66",
        spaceBefore=5 * mm,
        spaceAfter=3 * mm,
    )
    body_style = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=19,
        textColor="#1F2937",
        spaceAfter=3 * mm,
    )

    story = [
        Paragraph(escape(normalize_text(package.get("title", "微信公众号文章"))), title_style),
        Paragraph(escape(normalize_text(package.get("summary", ""))), lead_style),
    ]
    for section in package.get("sections", []):
        story.append(Paragraph(escape(normalize_text(section.get("heading", ""))), heading_style))
        for paragraph in section.get("paragraphs", []):
            story.append(Paragraph(escape(normalize_text(paragraph)), body_style))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("发布前检查", heading_style))
    for step in package.get("next_steps", []):
        story.append(Paragraph("• " + escape(normalize_text(step)), body_style))

    document = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=normalize_text(package.get("title", "微信公众号文章")),
        author="SciPoster",
    )
    document.build(story)
    return pdf_path


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Path to popular-article-package.json")
    parser.add_argument("--output-dir", required=True, help="Directory for exported files")
    args = parser.parse_args()

    package = json.loads(Path(args.input).read_text(encoding="utf-8"))
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = build_docx(package, output_dir)
    pdf_path = build_pdf(package, output_dir)

    print(json.dumps({"docx": str(docx_path), "pdf": str(pdf_path)}, ensure_ascii=False))


if __name__ == "__main__":
    main()
