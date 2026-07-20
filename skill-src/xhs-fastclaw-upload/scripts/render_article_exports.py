#!/usr/bin/env python
import argparse
import json
import os
from pathlib import Path
from typing import List, Tuple


try:
    from docx import Document
    from docx.shared import Inches
    from PIL import Image, ImageDraw, ImageFont  # type: ignore
except ImportError as exc:
    raise RuntimeError(
        "缺少 portable Python 依赖 python-docx、lxml 或 Pillow。请校验 runtime/python 并重新运行 deploy.ps1。"
    ) from exc


CARD_WIDTH = 1242
CARD_HEIGHT = 1660
THEME_GREEN = "#0E6B5B"
THEME_GREEN_LIGHT = "#EAF7F2"
THEME_TEXT = "#24333A"
THEME_MUTED = "#66747A"
THEME_LINE = "#D6E6E0"
THEME_BLACK = "#1C1F24"
THEME_SOFT = "#F7FCFA"


def load_font(size: int, bold: bool = False):
    override = os.environ.get("SCIPOSTER_FONT_DIR")
    font_root = Path(override) if override else Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    system_fonts = []
    if bold:
        system_fonts.extend(
            [
                font_root / "msyhbd.ttc",
                font_root / "simhei.ttf",
                font_root / "simsun.ttc",
            ]
        )
    system_fonts.extend(
        [
            font_root / "msyh.ttc",
            font_root / "simhei.ttf",
            font_root / "simsun.ttc",
        ]
    )

    for path in system_fonts:
        try:
            if path.exists():
                return ImageFont.truetype(str(path), size=size)
        except Exception:
            continue
    raise RuntimeError(
        "未找到可用的 Windows 中文字体。请安装 Microsoft YaHei（msyh.ttc/msyhbd.ttc）后重新部署。"
    )


def clean(text: str) -> str:
    return " ".join(str(text or "").replace("\r", " ").replace("\n", " ").split()).strip()


def split_tokens(text: str) -> List[str]:
    import re

    return re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9_\-/.%]+|[^\w\s]| +", text)


def wrap_text(draw, text: str, font, max_width: int, limit: int = None) -> List[str]:
    value = clean(text)
    if not value:
        return [""]
    tokens = split_tokens(value)
    lines = []
    current = ""
    for token in tokens:
        candidate = token if not current else current + token
        candidate_width = draw.textbbox((0, 0), candidate, font=font)[2]
        if current and candidate_width > max_width:
            lines.append(current.strip())
            current = token.strip()
        else:
            current = candidate
    if current.strip():
        lines.append(current.strip())
    return lines[:limit] if limit is not None else lines


def draw_multiline(draw, text: str, x: int, y: int, font, fill: str, max_width: int, line_gap: int, limit: int = None):
    cursor_y = y
    for line in wrap_text(draw, text, font, max_width, limit):
        draw.text((x, cursor_y), line, font=font, fill=fill)
        cursor_y += line_gap
    return cursor_y


def paste_paper_image(canvas: Image.Image, image_path: Path, box: Tuple[int, int, int, int]) -> bool:
    if not image_path.exists():
        return False
    try:
        image = Image.open(image_path).convert("RGB")
    except Exception:
        return False
    image.thumbnail((box[2] - box[0], box[3] - box[1]))
    x = box[0] + ((box[2] - box[0]) - image.width) // 2
    y = box[1] + ((box[3] - box[1]) - image.height) // 2
    canvas.paste(image, (x, y))
    return True


def render_visual_frame(canvas: Image.Image, image_path: Path, box: Tuple[int, int, int, int], border: str = "#D6E6E0") -> bool:
    frame = Image.new("RGB", (box[2] - box[0], box[3] - box[1]), "#F7FCFA")
    if not paste_paper_image(frame, image_path, (0, 0, frame.width, frame.height)):
        draw = ImageDraw.Draw(frame)
        draw.rounded_rectangle((2, 2, frame.width - 2, frame.height - 2), radius=24, outline=border, width=3)
        draw.text((frame.width // 2 - 42, frame.height // 2 - 10), "No image", fill=THEME_MUTED, font=load_font(24, bold=True))
    else:
        draw = ImageDraw.Draw(frame)
        draw.rounded_rectangle((2, 2, frame.width - 2, frame.height - 2), radius=24, outline=border, width=3)
    canvas.paste(frame, (box[0], box[1]))
    return True


def draw_header(draw, page_no: int, total_pages: int) -> None:
    draw.rectangle((0, 0, CARD_WIDTH, 12), fill=THEME_GREEN)
    draw.rounded_rectangle((1046, 36, 1188, 110), radius=34, fill="#98A19F")
    draw.text((1080, 52), f"{page_no}", font=load_font(32, bold=True), fill="#FFFFFF")
    draw.text((1122, 52), f"/{total_pages}", font=load_font(32, bold=True), fill="#FFFFFF")


def draw_footer(draw, page_no: int, total_pages: int) -> None:
    draw.line((0, CARD_HEIGHT - 132, CARD_WIDTH, CARD_HEIGHT - 132), fill="#E6ECE9", width=2)
    draw.text((60, CARD_HEIGHT - 96), "@SciPoster 小红书知识卡片", font=load_font(24), fill=THEME_MUTED)
    draw.text((1124, CARD_HEIGHT - 96), f"{page_no}/{total_pages}", font=load_font(24), fill=THEME_MUTED)


def render_cover(package: dict, image_dir: Path, output_path: Path, total_pages: int) -> None:
    canvas = Image.new("RGB", (CARD_WIDTH, CARD_HEIGHT), "#FFFFFF")
    draw = ImageDraw.Draw(canvas)
    draw_header(draw, 1, total_pages)

    cover_meta = package.get("cover_meta", {})
    asset_images = package.get("asset_images", []) or []
    hero_box = (0, 12, CARD_WIDTH, 598)

    if asset_images and paste_paper_image(canvas, image_dir / asset_images[0], hero_box):
        overlay = Image.new("RGBA", (CARD_WIDTH, 586), (255, 255, 255, 96))
        canvas.paste(overlay, (0, 12), overlay)
    else:
        draw.rectangle(hero_box, fill="#F3FBF7")
        draw.ellipse((40, 50, 420, 460), outline="#DAEFE7", width=3)
        draw.ellipse((836, 0, 1260, 410), outline="#DAEFE7", width=3)

    draw.rounded_rectangle((78, 76, 258, 152), radius=28, fill=THEME_GREEN)
    draw.text((112, 100), cover_meta.get("badge_left", "论文拆解"), font=load_font(30, bold=True), fill="#FFFFFF")
    draw.rounded_rectangle((286, 76, 664, 152), radius=28, outline=THEME_GREEN, width=2, fill="#F8FFFC")
    draw.text((334, 100), cover_meta.get("badge_right", "适合收藏复习"), font=load_font(30, bold=True), fill=THEME_GREEN)

    draw_multiline(draw, cover_meta.get("cover_title", package.get("chosen_title", "")), 78, 238, load_font(74, bold=True), THEME_GREEN, 1030, 88, 2)
    draw_multiline(draw, cover_meta.get("cover_subtitle", ""), 82, 396, load_font(44, bold=True), THEME_TEXT, 980, 54, 2)
    draw.rectangle((78, 500, 88, 602), fill=THEME_GREEN)
    draw_multiline(draw, cover_meta.get("summary", ""), 124, 512, load_font(28), THEME_TEXT, 948, 46, 4)

    draw.rounded_rectangle((70, 682, 1172, 908), radius=28, fill=THEME_GREEN_LIGHT)
    draw.text((96, 730), "论文核心结论", font=load_font(36, bold=True), fill=THEME_GREEN)
    draw_multiline(draw, package.get("summary", ""), 96, 792, load_font(30), THEME_TEXT, 1010, 50, 3)

    draw.rounded_rectangle((70, 968, 1172, 1260), radius=28, fill=THEME_SOFT, outline=THEME_LINE, width=2)
    draw.text((96, 1016), "这一组卡片会怎么展开", font=load_font(34, bold=True), fill=THEME_GREEN)
    cards = package.get("cards", []) or []
    y = 1078
    for card in cards[1:5]:
        y = draw_multiline(draw, "• " + clean(card.get("title", "")), 96, y, load_font(29), THEME_TEXT, 990, 46, 1) + 8

    draw.text((96, 1320), cover_meta.get("meta_line", ""), font=load_font(26, bold=True), fill=THEME_MUTED)
    draw_footer(draw, 1, total_pages)
    canvas.save(output_path)


def render_intro_card(card: dict, output_path: Path, page_no: int, total_pages: int) -> None:
    canvas = Image.new("RGB", (CARD_WIDTH, CARD_HEIGHT), "#FFFFFF")
    draw = ImageDraw.Draw(canvas)
    draw_header(draw, page_no, total_pages)
    draw.ellipse((30, 210, 380, 620), outline="#EAF5F1", width=3)
    draw.ellipse((860, 120, 1248, 532), outline="#EAF5F1", width=3)
    visual_image = card.get("visual_image")

    draw_multiline(draw, card.get("title", ""), 110, 154, load_font(70, bold=True), THEME_GREEN, 650, 84, 2)
    draw.line((110, 304, 240, 304), fill="#2CB494", width=8)
    draw_multiline(draw, card.get("lead", ""), 110, 360, load_font(30), THEME_TEXT, 650, 48, 2)

    y = 490
    for paragraph in card.get("paragraphs", []):
        y = draw_multiline(draw, paragraph, 110, y, load_font(29), THEME_TEXT, 650, 54, 4) + 34

    quote = clean(card.get("quote", ""))
    note = clean(card.get("note", ""))
    if quote:
        draw.rounded_rectangle((92, 1000, 1150, 1204), radius=18, fill=THEME_GREEN_LIGHT)
        draw.rectangle((92, 1000, 104, 1204), fill="#2CB494")
        draw_multiline(draw, f"“{quote}”", 132, 1046, load_font(32, bold=True), THEME_GREEN, 960, 52, 4)
    if note:
        draw_multiline(draw, note, 110, 1270, load_font(26), THEME_MUTED, 650, 42, 3)

    if visual_image:
        render_visual_frame(canvas, Path(visual_image), (792, 166, 1160, 732))

    draw_footer(draw, page_no, total_pages)
    canvas.save(output_path)


def render_bullet_card(card: dict, output_path: Path, page_no: int, total_pages: int) -> None:
    canvas = Image.new("RGB", (CARD_WIDTH, CARD_HEIGHT), "#FFFFFF")
    draw = ImageDraw.Draw(canvas)
    draw_header(draw, page_no, total_pages)
    visual_image = card.get("visual_image")

    draw_multiline(draw, card.get("title", ""), 88, 120, load_font(60, bold=True), THEME_GREEN, 650, 76, 2)
    draw.line((92, 236, 248, 236), fill="#2CB494", width=8)
    draw_multiline(draw, card.get("lead", ""), 92, 294, load_font(28), THEME_TEXT, 650, 48, 3)

    y = 436
    for bullet in card.get("bullets", []):
        term = clean(bullet.get("term", ""))
        desc = clean(bullet.get("desc", ""))
        draw.rounded_rectangle((92, y + 6, 190, y + 54), radius=18, fill=THEME_GREEN_LIGHT)
        draw.text((112, y + 12), term, font=load_font(24, bold=True), fill=THEME_GREEN)
        y = draw_multiline(draw, desc, 220, y, load_font(29), THEME_TEXT, 650, 48, 5) + 34

    if visual_image:
        render_visual_frame(canvas, Path(visual_image), (786, 918, 1160, 1380))

    draw_footer(draw, page_no, total_pages)
    canvas.save(output_path)


def render_insight_card(card: dict, output_path: Path, page_no: int, total_pages: int) -> None:
    canvas = Image.new("RGB", (CARD_WIDTH, CARD_HEIGHT), "#FFFFFF")
    draw = ImageDraw.Draw(canvas)
    draw_header(draw, page_no, total_pages)
    visual_image = card.get("visual_image")

    draw_multiline(draw, card.get("title", ""), 92, 120, load_font(62, bold=True), THEME_GREEN, 650, 76, 2)
    draw.line((92, 232, 244, 232), fill="#2CB494", width=8)
    draw_multiline(draw, card.get("lead", ""), 92, 288, load_font(28), THEME_TEXT, 650, 46, 2)

    y = 420
    for paragraph in card.get("paragraphs", []):
        y = draw_multiline(draw, paragraph, 92, y, load_font(30), THEME_TEXT, 1040, 54, 4) + 32

    draw.rounded_rectangle((86, 980, 1156, 1262), radius=20, fill=THEME_GREEN_LIGHT)
    draw.rectangle((86, 980, 96, 1262), fill="#2CB494")
    draw_multiline(draw, card.get("callout", ""), 128, 1034, load_font(31, bold=True), THEME_GREEN, 650, 54, 5)

    if visual_image:
        render_visual_frame(canvas, Path(visual_image), (786, 164, 1160, 698))

    draw_footer(draw, page_no, total_pages)
    canvas.save(output_path)


def render_ending_card(card: dict, output_path: Path, page_no: int, total_pages: int) -> None:
    canvas = Image.new("RGB", (CARD_WIDTH, CARD_HEIGHT), THEME_BLACK)
    draw = ImageDraw.Draw(canvas)
    visual_image = card.get("visual_image")
    draw_multiline(draw, card.get("title", ""), 90, 118, load_font(58, bold=True), "#FFFFFF", 650, 72, 2)
    draw_multiline(draw, card.get("lead", ""), 92, 234, load_font(28), "#E9ECEE", 650, 42, 2)

    y = 366
    for paragraph in card.get("paragraphs", []):
        y = draw_multiline(draw, paragraph, 92, y, load_font(31), "#F5F7F8", 650, 58, 5) + 52

    dots_x = 558
    for idx in range(total_pages):
        color = "#F04A5B" if idx + 1 == page_no else "#D9D9D9"
        draw.ellipse((dots_x + idx * 26, 72, dots_x + idx * 26 + 12, 84), fill=color)

    if visual_image:
        render_visual_frame(canvas, Path(visual_image), (788, 342, 1160, 804), border="#56606A")

    draw_footer(draw, page_no, total_pages)
    canvas.save(output_path)


def build_docx(package: dict, image_files: List[Path], output_path: Path) -> None:
    document = Document()
    document.add_heading(package.get("chosen_title", "小红书图文方案"), level=0)
    document.add_paragraph("这是适合直接发布的小红书图文方案总说明。")

    document.add_heading("可直接复制内容", level=1)
    document.add_paragraph(f"标题：{package.get('chosen_title', '')}")
    document.add_paragraph(f"正文：{package.get('post_copy', package.get('summary', ''))}")
    document.add_paragraph(f"标签：{' '.join(package.get('hashtags', []))}")
    document.add_paragraph(f"置顶评论：{package.get('pinned_comment', '欢迎留言交流你最关心的论文内容。')}")

    document.add_heading("卡片顺序", level=1)
    for idx, image_path in enumerate(image_files, start=1):
        document.add_paragraph(f"{idx}. {image_path.name}")

    document.add_heading("卡片预览缩略图", level=1)
    for image_path in image_files:
        document.add_paragraph(image_path.name)
        document.add_picture(str(image_path), width=Inches(4.8))

    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Path to xiaohongshu-package.json")
    parser.add_argument("--output-dir", required=True, help="Directory for exported files")
    args = parser.parse_args()

    package = json.loads(Path(args.input).read_text(encoding="utf-8"))
    output_dir = Path(args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    image_dir = output_dir / "xiaohongshu-images"
    image_dir.mkdir(parents=True, exist_ok=True)

    cards = package.get("cards", []) or []
    total_pages = len(cards)
    generated_images = []

    for idx, card in enumerate(cards, start=1):
        out_path = image_dir / f"card{idx}.png"
        kind = card.get("kind")
        if kind == "cover":
            render_cover(package, output_dir / "extracted-paper-assets", out_path, total_pages)
        elif kind == "intro":
            render_intro_card(card, out_path, idx, total_pages)
        elif kind == "bullet_explain":
            render_bullet_card(card, out_path, idx, total_pages)
        elif kind == "insight":
            render_insight_card(card, out_path, idx, total_pages)
        else:
            render_ending_card(card, out_path, idx, total_pages)
        generated_images.append(out_path)

    package["generated_images"] = [path.name for path in generated_images]
    Path(args.input).write_text(json.dumps(package, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    docx_path = output_dir / "xiaohongshu-package.docx"
    build_docx(package, generated_images, docx_path)

    print(json.dumps({"docx": str(docx_path), "images": [str(path) for path in generated_images]}, ensure_ascii=False))


if __name__ == "__main__":
    main()
