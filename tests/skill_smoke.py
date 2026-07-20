#!/usr/bin/env python
"""End-to-end smoke tests for version-controlled SciPoster skill sources."""

from __future__ import annotations

import argparse
import json
import os
import subprocess
import sys
import tempfile
import zipfile
from html.parser import HTMLParser
from pathlib import Path
from xml.etree import ElementTree

from PIL import Image
from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


PAPER_TEXT = """A Minimal SciPoster Paper

Abstract
This study evaluates a reproducible local publishing workflow.

Methods
We process one small document with deterministic scripts.

Results
All expected artifacts are generated and remain parseable.

Conclusion
Portable runtimes make the deployment easier to verify.
"""

RUNNERS = {
    "academic-poster-fastclaw-upload": "run_academic_poster_pipeline.py",
    "slides-fastclaw-upload": "run_academic_slides_pipeline.py",
    "popular-article-fastclaw-upload": "run_popular_article_pipeline.py",
    "xhs-fastclaw-upload": "run_xhs_pipeline.py",
}


class StrictHTMLParser(HTMLParser):
    pass


def create_input(workspace: Path, input_format: str) -> None:
    if input_format == "txt":
        (workspace / "paper.txt").write_text(PAPER_TEXT, encoding="utf-8")
    elif input_format == "docx":
        document = Document()
        document.add_heading("A Minimal SciPoster Paper", level=0)
        for block in PAPER_TEXT.split("\n\n")[1:]:
            document.add_paragraph(block)
        document.save(workspace / "paper.docx")
    elif input_format == "pdf":
        output = workspace / "paper.pdf"
        pdf = canvas.Canvas(str(output), pagesize=A4)
        text = pdf.beginText(50, 790)
        text.setFont("Helvetica", 11)
        for line in PAPER_TEXT.splitlines():
            text.textLine(line)
        pdf.drawText(text)
        pdf.save()
    else:
        raise ValueError(input_format)


def validate_zip(path: Path) -> None:
    if not zipfile.is_zipfile(path):
        raise AssertionError(f"Expected a valid ZIP container: {path}")
    with zipfile.ZipFile(path) as archive:
        if "[Content_Types].xml" not in archive.namelist():
            raise AssertionError(f"Office document lacks [Content_Types].xml: {path}")


def validate_png(path: Path) -> None:
    with Image.open(path) as image:
        image.verify()


def validate_html(path: Path) -> None:
    parser = StrictHTMLParser()
    parser.feed(path.read_text(encoding="utf-8"))
    parser.close()


def validate_json(path: Path) -> None:
    json.loads(path.read_text(encoding="utf-8"))


def validate_pdf(path: Path) -> None:
    if not path.read_bytes().startswith(b"%PDF-"):
        raise AssertionError(f"Not a PDF file: {path}")
    reader = PdfReader(str(path))
    if not reader.pages:
        raise AssertionError(f"PDF has no pages: {path}")


def validate_outputs(skill: str, output: Path) -> None:
    validate_json(output / "final-response.json")
    if skill == "academic-poster-fastclaw-upload":
        validate_zip(output / "academic-poster.pptx")
        validate_png(output / "academic-poster.png")
        ElementTree.parse(output / "academic-poster.svg")
        validate_html(output / "poster-preview.html")
        validate_json(output / "poster-package.json")
    elif skill == "slides-fastclaw-upload":
        validate_zip(output / "slides.pptx")
        validate_png(output / "slides-preview.png")
        for page in (output / "slides-pages").glob("*.png"):
            validate_png(page)
        validate_html(output / "slides-preview.html")
        validate_json(output / "slides-package.json")
    elif skill == "popular-article-fastclaw-upload":
        validate_zip(output / "popular-article.docx")
        validate_zip(output / "wechat-cover.pptx")
        validate_png(output / "wechat-cover.png")
        validate_pdf(output / "popular-article.pdf")
        validate_html(output / "popular-article-preview.html")
        validate_json(output / "popular-article-package.json")
        if (output / "popular-article.doc").exists():
            raise AssertionError("Deprecated fake popular-article.doc was generated")
    elif skill == "xhs-fastclaw-upload":
        validate_zip(output / "xiaohongshu-package.docx")
        cards = sorted((output / "xiaohongshu-images").glob("*.png"))
        if not cards:
            raise AssertionError("XHS pipeline generated no cards")
        for card in cards:
            validate_png(card)
        validate_html(output / "xiaohongshu-preview.html")
        validate_json(output / "xiaohongshu-package.json")


def run_pipeline(repo_root: Path, skill: str, input_format: str, node: str) -> None:
    with tempfile.TemporaryDirectory(prefix=f"sciposter-{skill}-{input_format}-") as temporary:
        workspace = Path(temporary)
        output = workspace / "output"
        create_input(workspace, input_format)
        runner = repo_root / "skill-src" / skill / "scripts" / RUNNERS[skill]
        command = [
            sys.executable,
            str(runner),
            "--workspace",
            str(workspace),
            "--output-dir",
            str(output),
            "--python",
            sys.executable,
        ]
        if skill in {"academic-poster-fastclaw-upload", "slides-fastclaw-upload"}:
            command.extend(["--node", node])
        subprocess.run(command, check=True, env=os.environ.copy())
        validate_outputs(skill, output)
        print(f"PASS {skill} {input_format}")


def run_missing_input(repo_root: Path, skill: str, node: str) -> None:
    with tempfile.TemporaryDirectory(prefix=f"sciposter-{skill}-missing-") as temporary:
        workspace = Path(temporary)
        output = workspace / "output"
        runner = repo_root / "skill-src" / skill / "scripts" / RUNNERS[skill]
        command = [
            sys.executable,
            str(runner),
            "--workspace",
            str(workspace),
            "--output-dir",
            str(output),
            "--python",
            sys.executable,
        ]
        if skill in {"academic-poster-fastclaw-upload", "slides-fastclaw-upload"}:
            command.extend(["--node", node])
        subprocess.run(command, check=True, env=os.environ.copy())
        report = output / "missing-upload-report.json"
        validate_json(report)
        if (output / "final-response.json").exists():
            raise AssertionError(f"Missing-input path incorrectly reported success: {skill}")
        print(f"PASS {skill} missing-input")


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser()
    parser.add_argument("--skill", choices=["all", *RUNNERS], default="all")
    parser.add_argument("--input-format", choices=["all", "txt", "docx", "pdf"], default="txt")
    parser.add_argument("--missing-input", action="store_true")
    parser.add_argument("--missing-only", action="store_true")
    parser.add_argument("--node", default=os.environ.get("SCIPOSTER_NODE", "node"))
    args = parser.parse_args()

    skills = list(RUNNERS) if args.skill == "all" else [args.skill]
    formats = ["txt", "docx", "pdf"] if args.input_format == "all" else [args.input_format]
    for skill in skills:
        if not args.missing_only:
            for input_format in formats:
                run_pipeline(repo_root, skill, input_format, args.node)
        if args.missing_input or args.missing_only:
            run_missing_input(repo_root, skill, args.node)


if __name__ == "__main__":
    main()
